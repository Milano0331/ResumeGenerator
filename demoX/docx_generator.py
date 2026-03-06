import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_font(run, font_name="Microsoft YaHei", font_size=10, bold=False, italic=False, color=None):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def create_resume_docx(markdown_content, output_file):
    """
    Convert markdown content to a DOCX resume.
    """
    document = Document()
    
    # Set margins (A4 standard: 1 inch = 2.54 cm)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse content line by line
    lines = markdown_content.split('\n')
    
    # Basic parsing logic
    # Line 0: Name (e.g., "# Name")
    # Line 2: Contact
    
    name_line = "Resume"
    contact_line = ""
    start_index = 0
    
    if len(lines) > 0:
        name_line = lines[0].strip().replace('# ', '').replace('#', '')
        start_index += 1
    
    # Try to find contact info (often lines 1-3)
    for i in range(1, min(5, len(lines))):
        line = lines[i].strip()
        if line and not line.startswith('#') and not line.startswith('-'):
            contact_line = line
            start_index = i + 1
            break
            
    # --- Header Section ---
    # Name
    h1 = document.add_heading(level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(name_line)
    set_font(run, font_name="Microsoft YaHei", font_size=18, bold=True, color=RGBColor(0, 0, 0))
    
    # Contact
    if contact_line:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_line)
        set_font(run, font_name="Microsoft YaHei", font_size=10, color=RGBColor(80, 80, 80)) # Dark Gray
        
    # Horizontal Line
    document.add_paragraph().add_run().add_break() 

    # --- Body Content ---
    for line in lines[start_index:]:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('## '):
            # Heading 2
            text = line.replace('## ', '')
            h2 = document.add_heading(level=2)
            run = h2.add_run(text)
            set_font(run, font_name="Microsoft YaHei", font_size=14, bold=True, color=RGBColor(0, 0, 0))
            # Add some space before
            h2.paragraph_format.space_before = Pt(12)
            h2.paragraph_format.space_after = Pt(6)
            
        elif line.startswith('### '):
             # Heading 3 (Treat as subheading)
             text = line.replace('### ', '')
             h3 = document.add_heading(level=3)
             run = h3.add_run(text)
             set_font(run, font_name="Microsoft YaHei", font_size=12, bold=True, color=RGBColor(0, 0, 0))
             
        elif line.startswith('---'):
            # Horizontal Rule (Simulated with empty paragraph for now, or could use border)
            document.add_paragraph().add_run("________________________________________________________________________________")
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet Point
            text = line[2:]
            p = document.add_paragraph(style='List Bullet')
            
            # Simple bold/italic parsing
            # Split by ** for bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    content = part[2:-2]
                    run = p.add_run(content)
                    set_font(run, font_name="Microsoft YaHei", font_size=10, bold=True)
                else:
                    # Check for italic *
                    sub_parts = re.split(r'(\*.*?\*)', part)
                    for sub_part in sub_parts:
                        if sub_part.startswith('*') and sub_part.endswith('*'):
                            content = sub_part[1:-1]
                            run = p.add_run(content)
                            set_font(run, font_name="Microsoft YaHei", font_size=10, italic=True)
                        else:
                            run = p.add_run(sub_part)
                            set_font(run, font_name="Microsoft YaHei", font_size=10)
                            
        else:
            # Normal Paragraph
            p = document.add_paragraph()
            
            # Simple bold/italic parsing
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    content = part[2:-2]
                    run = p.add_run(content)
                    set_font(run, font_name="Microsoft YaHei", font_size=10, bold=True)
                else:
                    # Check for italic *
                    sub_parts = re.split(r'(\*.*?\*)', part)
                    for sub_part in sub_parts:
                        if sub_part.startswith('*') and sub_part.endswith('*'):
                            content = sub_part[1:-1]
                            run = p.add_run(content)
                            set_font(run, font_name="Microsoft YaHei", font_size=10, italic=True)
                        else:
                            run = p.add_run(sub_part)
                            set_font(run, font_name="Microsoft YaHei", font_size=10)

    # Save
    try:
        document.save(output_file)
        print(f"Successfully generated DOCX: {output_file}")
        return True
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return False
